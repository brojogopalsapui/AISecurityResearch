from pathlib import Path
import re
from calendar import month_name
from html import escape

from bs4 import BeautifulSoup, NavigableString
from docx import Document

# ===== USER CONFIG =====
from pathlib import Path

if "__file__" in globals():
    BASE_DIR = Path(__file__).resolve().parent
else:
    BASE_DIR = Path.cwd()

# If running from dark folder, go one level up
if BASE_DIR.name == "dark":
    PROJECT_ROOT = BASE_DIR.parent
else:
    PROJECT_ROOT = BASE_DIR

WEEKLY_INPUTS_DIR = PROJECT_ROOT / 'weekly-inputs'
POSTS_DIR = PROJECT_ROOT / 'dark' / 'posts'
BATCH_PROCESS_ALL_DOCX = True
AUTO_PICK_LATEST_DOCX = False
MANUAL_DOCX_PATH = WEEKLY_INPUTS_DIR / '2026-03-stack-safeguard-pipelines.docx'
INPUT_ONGOING_HTML = PROJECT_ROOT / 'dark' / 'ongoing-work.html'
INPUT_INDEX_HTML = PROJECT_ROOT / 'dark' / 'index.html'
OUTPUT_ONGOING_HTML = INPUT_ONGOING_HTML
OUTPUT_INDEX_HTML = INPUT_INDEX_HTML
AUTO_GENERATE_FULL_POST = True
OUTPUT_POSTS_DIR = POSTS_DIR
REPLACE_DUPLICATE_POST_ID = True
UPDATE_HOME_FLOATING = True
UPDATE_HOME_SLIDER = True
HOME_SLIDER_MAX_CARDS = 6
FORCE_POST_ID_FROM_FILENAME = True
AUTO_USE_FILENAME_BASED_FULL_POST_LINK = True

TEMPLATE_DOCX_HINTS = {
    'weekly-research-watch-template.docx',
    'template.docx',
    'research-watch-template.docx',
}
ALLOWED_CATEGORIES = ('academic', 'industry', 'ecosystem')
CATEGORY_LABELS = {
    'academic': 'Academic Signals',
    'industry': 'Industry & Innovation',
    'ecosystem': 'Companies & Releases',
}
CATEGORY_PILL_LABELS = {
    'academic': 'Academic',
    'industry': 'Industry',
    'ecosystem': 'Ecosystem',
}
CATEGORY_META_LABELS = {
    'academic': 'Academic stream',
    'industry': 'Industry stream',
    'ecosystem': 'Ecosystem stream',
}
SECTION_TARGETS = {
    'academic': 'academic-stream-list',
    'industry': 'industry-stream-list',
    'ecosystem': 'ecosystem-stream-list',
}

PREFERRED_FILENAME_RE = re.compile(r'^(\d{2})-(\d{2})-(\d{4})-(\d{6})-([a-z0-9][a-z0-9-]*)$', re.IGNORECASE)
YMD_FILENAME_RE = re.compile(r'^(\d{4})-(\d{2})-(\d{2})-([a-z0-9][a-z0-9-]*)$', re.IGNORECASE)
LEGACY_FILENAME_RE = re.compile(r'^(\d{4})-(\d{2})-(.+)$', re.IGNORECASE)


def _clean(value):
    if value is None:
        return ''
    if isinstance(value, str):
        return value.strip()
    return str(value).strip()


def parse_filename_info(stem: str):
    m = PREFERRED_FILENAME_RE.match(stem)
    if m:
        day, month, year = int(m.group(1)), int(m.group(2)), int(m.group(3))
        hhmmss = m.group(4)
        hour, minute, second = int(hhmmss[0:2]), int(hhmmss[2:4]), int(hhmmss[4:6])
        slug = m.group(5).lower()
        if 1 <= month <= 12 and 1 <= day <= 31:
            return {
                'style': 'preferred',
                'year': year,
                'month': month,
                'day': day,
                'hour': hour,
                'minute': minute,
                'second': second,
                'slug': slug,
                'display_date': f'{day:02d}-{month:02d}-{year:04d}',
                'month_year': f'{month_name[month]} {year}',
            }
    m = YMD_FILENAME_RE.match(stem)
    if m:
        year, month, day = int(m.group(1)), int(m.group(2)), int(m.group(3))
        slug = m.group(4).lower()
        if 1 <= month <= 12 and 1 <= day <= 31:
            return {
                'style': 'ymd',
                'year': year,
                'month': month,
                'day': day,
                'hour': 0,
                'minute': 0,
                'second': 0,
                'slug': slug,
                'display_date': f'{day:02d}-{month:02d}-{year:04d}',
                'month_year': f'{month_name[month]} {year}',
            }
    m = LEGACY_FILENAME_RE.match(stem)
    if m:
        year, month = int(m.group(1)), int(m.group(2))
        slug = m.group(3).lower()
        if 1 <= month <= 12:
            return {
                'style': 'legacy',
                'year': year,
                'month': month,
                'day': 0,
                'hour': 0,
                'minute': 0,
                'second': 0,
                'slug': slug,
                'display_date': f'{month:02d}-{year:04d}',
                'month_year': f'{month_name[month]} {year}',
            }
    return None


def stem_to_month_year(stem: str) -> str:
    info = parse_filename_info(stem)
    return info['month_year'] if info else stem


def stem_to_display_date(stem: str) -> str:
    info = parse_filename_info(stem)
    return info['display_date'] if info else ''


def format_trending_title(stem: str, title: str) -> str:
    display_date = stem_to_display_date(stem)
    return f'{display_date} | {title}' if display_date else title


def _docx_sort_key(path: Path):
    info = parse_filename_info(path.stem)
    if info:
        return (0, info['year'], info['month'], info['day'], info['hour'], info['minute'], info['second'], info['slug'])
    return (1, path.stat().st_mtime, path.stem.lower())


def _filename_guidance(name: str) -> str:
    return (
        f"Filename '{name}' does not match the preferred DD-MM-YYYY-HHMMSS-topic-name.docx format. "
        "Use zero-padded date and HHMMSS so the latest note is selected correctly even when multiple notes are added on the same day."
    )


def list_processable_docx_files(folder: Path):
    hints = {name.lower() for name in TEMPLATE_DOCX_HINTS}
    candidates = []
    for f in folder.glob('*.docx'):
        if f.name.lower() in hints or f.name.startswith('~$'):
            continue
        candidates.append(f)
    if not candidates:
        raise FileNotFoundError(f'No usable DOCX files found in {folder}.')
    return sorted(candidates, key=_docx_sort_key)


def pick_latest_docx(folder: Path) -> Path:
    return list_processable_docx_files(folder)[-1]


def parse_weekly_docx(path: Path):
    doc = Document(str(path))
    data = {'metadata': {}}
    if doc.tables:
        table = doc.tables[0]
        for row in table.rows[1:]:
            if len(row.cells) < 2:
                continue
            key = _clean(row.cells[0].text)
            val = _clean(row.cells[1].text)
            if key:
                data['metadata'][key] = val
    wanted_headings = [
        'Full Note Paragraph 1',
        'Full Note Paragraph 2',
        'What Is Changing Technically',
        'What Reviewers Should Notice',
        'Current Research Tension',
    ]
    paragraphs = doc.paragraphs
    for i, p in enumerate(paragraphs):
        heading = _clean(p.text)
        if heading in wanted_headings:
            j = i + 1
            content = []
            while j < len(paragraphs):
                t = _clean(paragraphs[j].text)
                if t in wanted_headings or t in {'Metadata', 'How the notebook uses this DOCX'}:
                    break
                if t:
                    content.append(t)
                j += 1
            data[heading] = content if heading in {'What Is Changing Technically', 'What Reviewers Should Notice'} else '\n'.join(content)
    return data




def is_placeholder_docx(data) -> bool:
    md = data.get('metadata', {})
    checks = [
        _clean(md.get('Title')).lower(),
        _clean(md.get('Meta Line')).lower(),
        _clean(md.get('Preview')).lower(),
        _clean(md.get('Related Static Page (optional)')).lower(),
        _clean(md.get('External Link 1 URL (optional)')).lower(),
        _clean(md.get('External Link 2 URL (optional)')).lower(),
    ]
    placeholder_markers = [
        'your new weekly research note title',
        'research watch • your topic',
        'write a short 2–3 line preview',
        'write a short 2-3 line preview',
        'ai-security/related-page.html',
        'https://example.com',
    ]
    return any(marker in text for text in checks for marker in placeholder_markers if text)

def normalize_category(raw_category: str, warnings: list[str]) -> str:
    category = _clean(raw_category).lower()
    if category not in ALLOWED_CATEGORIES:
        if category:
            warnings.append(f"Category '{raw_category}' is invalid. Allowed values: {', '.join(ALLOWED_CATEGORIES)}. Falling back to 'academic'.")
        else:
            warnings.append("Category was blank in the DOCX. Falling back to 'academic'.")
        category = 'academic'
    return category


def normalize_docx_data(data, docx_stem: str, relative_full_post_link: str):
    md = data['metadata']
    warnings = []
    if not PREFERRED_FILENAME_RE.match(docx_stem):
        warnings.append(_filename_guidance(f'{docx_stem}.docx'))
    if FORCE_POST_ID_FROM_FILENAME:
        existing_post_id = _clean(md.get('Post ID'))
        if existing_post_id and existing_post_id != docx_stem:
            warnings.append(f"Overriding DOCX Post ID '{existing_post_id}' with filename-based Post ID '{docx_stem}' for consistency.")
        md['Post ID'] = docx_stem
    elif not _clean(md.get('Post ID')):
        md['Post ID'] = docx_stem
    if AUTO_USE_FILENAME_BASED_FULL_POST_LINK:
        existing_link = _clean(md.get('Full Post Link (optional)'))
        if existing_link and existing_link != relative_full_post_link:
            warnings.append(f"Overriding DOCX Full Post Link '{existing_link}' with filename-based link '{relative_full_post_link}' for consistency.")
        md['Full Post Link (optional)'] = relative_full_post_link
    md['Category'] = normalize_category(md.get('Category'), warnings)
    if not _clean(md.get('Title')):
        md['Title'] = docx_stem.replace('-', ' ').title()
        warnings.append('Title was blank in the DOCX, so a title was inferred from the filename.')
    if not _clean(md.get('Meta Line')):
        md['Meta Line'] = f'Research Watch • {stem_to_month_year(docx_stem)} signal'
        warnings.append('Meta Line was blank in the DOCX, so one was inferred from the filename.')
    if not _clean(md.get('Preview')):
        p1 = _clean(data.get('Full Note Paragraph 1'))
        md['Preview'] = (p1[:220].rsplit(' ', 1)[0].rstrip('.,;: ') + '...') if p1 else md['Title']
        warnings.append('Preview was blank in the DOCX, so one was inferred from the first full-note paragraph.')
    return warnings


def _is_valid_href(href: str) -> bool:
    href = _clean(href)
    if not href:
        return False
    if href.startswith(('http://', 'https://', 'mailto:', '#')):
        return True
    if ' ' in href:
        return False
    return '/' in href or href.endswith('.html') or href.endswith('.pdf')


def add_link_html(links, href, label):
    href = _clean(href)
    label = _clean(label)
    if not _is_valid_href(href) or not label:
        return
    if href.startswith(('http://', 'https://')):
        links.append(f'<a href="{escape(href, quote=True)}" rel="noopener noreferrer" target="_blank">{escape(label)}</a>')
    else:
        links.append(f'<a href="{escape(href, quote=True)}">{escape(label)}</a>')


def build_home_watch_card(post_id, title, preview):
    return f'''
<a class="watch-card" href="ongoing-work.html#{escape(post_id, quote=True)}">
  <div class="watch-card-label">Latest research note</div>
  <h3>{escape(title)}</h3>
  <p>{escape(preview)}</p>
  <span>Open note -></span>
</a>'''.strip()


def build_watch_article(data):
    md = data['metadata']
    category = _clean(md.get('Category')).lower()
    post_id = _clean(md.get('Post ID'))
    title = _clean(md.get('Title')) or 'Untitled research note'
    display_title = format_trending_title(post_id, title)
    p1 = _clean(data.get('Full Note Paragraph 1'))
    p2 = _clean(data.get('Full Note Paragraph 2'))
    preview = _clean(md.get('Preview')) or ((p1[:220].rsplit(' ', 1)[0] + '...') if p1 else title)
    meta_line = _clean(md.get('Meta Line')) or f'Research Watch • {stem_to_month_year(post_id)} signal'
    stream_label = CATEGORY_META_LABELS.get(category, category.title())
    if stream_label.lower() not in meta_line.lower():
        meta_line = f'{meta_line} • {stream_label}'
    category_pill = CATEGORY_PILL_LABELS.get(category, category.title())
    tech_list = data.get('What Is Changing Technically', [])
    reviewer_list = data.get('What Reviewers Should Notice', [])
    tension = _clean(data.get('Current Research Tension'))
    links = []
    add_link_html(links, md.get('Full Post Link (optional)'), 'Read full post')
    add_link_html(links, md.get('Related Static Page (optional)'), md.get('Related Static Page Label (optional)'))
    add_link_html(links, md.get('External Link 1 URL (optional)'), md.get('External Link 1 Label (optional)'))
    add_link_html(links, md.get('External Link 2 URL (optional)'), md.get('External Link 2 Label (optional)'))
    tech_items = '\n'.join(f'<li>{escape(_clean(item))}</li>' for item in tech_list if _clean(item))
    reviewer_items = '\n'.join(f'<li>{escape(_clean(item))}</li>' for item in reviewer_list if _clean(item))
    links_html = ' '.join(links)
    article_html = f'''
<article class="watch-note accordion" data-category="{escape(category, quote=True)}" id="{escape(post_id, quote=True)}">
  <button aria-expanded="false" class="accordion-trigger" type="button">
    <span class="accordion-meta"><span class="category-pill category-pill-{escape(category, quote=True)}">{escape(category_pill)}</span><span class="meta-line-text">{escape(meta_line)}</span></span>
    <span class="accordion-title">{escape(display_title)}</span>
    <span class="accordion-preview">{escape(preview)}</span>
    <span class="accordion-cta">Read full note</span>
    <span aria-hidden="true" class="accordion-icon"></span>
  </button>
  <div aria-hidden="true" class="accordion-panel">
    <div class="accordion-panel-inner">
      <p>{escape(p1)}</p>
      <p>{escape(p2)}</p>
      <div class="watch-columns">
        <div>
          <h4>What is changing technically</h4>
          <ul>{tech_items}</ul>
        </div>
        <div>
          <h4>What reviewers should notice</h4>
          <ul>{reviewer_items}</ul>
        </div>
      </div>
      <p class="watch-tension"><strong>Current research tension.</strong> {escape(tension)}</p>
      <div class="watch-links">{links_html}</div>
    </div>
  </div>
</article>
'''.strip()
    return post_id, title, preview, category, article_html


def build_full_post_html(data):
    md = data['metadata']
    category = _clean(md.get('Category')).lower()
    title = _clean(md.get('Title'))
    meta_line = _clean(md.get('Meta Line'))
    p1 = _clean(data.get('Full Note Paragraph 1'))
    p2 = _clean(data.get('Full Note Paragraph 2'))
    tech_list = data.get('What Is Changing Technically', [])
    reviewer_list = data.get('What Reviewers Should Notice', [])
    tension = _clean(data.get('Current Research Tension'))
    links = []
    add_link_html(links, md.get('Related Static Page (optional)'), md.get('Related Static Page Label (optional)'))
    add_link_html(links, md.get('External Link 1 URL (optional)'), md.get('External Link 1 Label (optional)'))
    add_link_html(links, md.get('External Link 2 URL (optional)'), md.get('External Link 2 Label (optional)'))
    links_html = ' '.join(links)
    tech_items = '\n'.join(f'<li>{escape(_clean(item))}</li>' for item in tech_list if _clean(item))
    reviewer_items = '\n'.join(f'<li>{escape(_clean(item))}</li>' for item in reviewer_list if _clean(item))
    stream_badge = CATEGORY_LABELS.get(category, category.title())
    return f'''<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <meta content="width=device-width, initial-scale=1.0" name="viewport" />
  <title>{escape(title)} | AI Security Research</title>
  <link href="../assets/css/style.css" rel="stylesheet" />
</head>
<body>
  <main class="section" style="padding-top: 8rem;">
    <div class="container" style="max-width: 900px;">
      <a href="../ongoing-work.html" style="display:inline-block;margin-bottom:1rem;"><- Back to Trending Topics</a>
      <div class="section-label">{escape(stream_badge)}</div>
      <h1>{escape(title)}</h1>
      <p class="section-lead">{escape(meta_line)}</p>
      <p>{escape(p1)}</p>
      <p>{escape(p2)}</p>
      <h2>What is changing technically</h2>
      <ul>{tech_items}</ul>
      <h2>What reviewers should notice</h2>
      <ul>{reviewer_items}</ul>
      <p><strong>Current research tension.</strong> {escape(tension)}</p>
      <div class="watch-links">{links_html}</div>
    </div>
  </main>
  <script src="../assets/js/main.js"></script>
</body>
</html>'''.strip()


def detect_doctype(html_text: str) -> str:
    m = re.match(r'\s*(<!DOCTYPE[^>]+>)', html_text, flags=re.IGNORECASE)
    return m.group(1) if m else '<!DOCTYPE html>'


def validate_paths(docx_path, ongoing_input_path, ongoing_output_path, index_input_path, index_output_path):
    warnings = []
    if not docx_path.exists():
        raise FileNotFoundError(f'DOCX file not found: {docx_path.resolve()}')
    if not ongoing_input_path.exists():
        raise FileNotFoundError(f'Input ongoing-work.html not found: {ongoing_input_path.resolve()}')
    if not index_input_path.exists():
        raise FileNotFoundError(f'Input index.html not found: {index_input_path.resolve()}')
    if docx_path.name.lower() in {name.lower() for name in TEMPLATE_DOCX_HINTS}:
        warnings.append('DOCX_PATH looks like a blank template file. Use a filled weekly DOCX instead.')
    if not PREFERRED_FILENAME_RE.match(docx_path.stem):
        warnings.append(_filename_guidance(docx_path.name))
    if ongoing_input_path.resolve() != ongoing_output_path.resolve():
        warnings.append('OUTPUT_ONGOING_HTML differs from INPUT_ONGOING_HTML. Best default: update in place.')
    if index_input_path.resolve() != index_output_path.resolve():
        warnings.append('OUTPUT_INDEX_HTML differs from INPUT_INDEX_HTML. Best default: update in place.')
    return warnings


def _find_stream_container(soup: BeautifulSoup, category: str):
    target_id = SECTION_TARGETS[category]
    container = soup.select_one(f'#{target_id}')
    if container is not None:
        return container
    watch_stack = soup.select_one('div.watch-stack')
    if watch_stack is not None:
        return watch_stack
    raise ValueError('Could not find a supported trending-topics container in ongoing-work.html')


def update_ongoing_work_html(input_html_path, output_html_path, article_id, article_html, category, replace_duplicate=True):
    html_text = input_html_path.read_text(encoding='utf-8')
    doctype = detect_doctype(html_text)
    soup = BeautifulSoup(html_text, 'html.parser')
    if replace_duplicate:
        existing = soup.find('article', {'id': article_id})
        if existing:
            existing.decompose()
    target = _find_stream_container(soup, category)
    fragment = BeautifulSoup(article_html, 'html.parser')
    new_article = fragment.find('article')
    first_existing = target.find('article')
    if first_existing:
        first_existing.insert_before('\n')
        first_existing.insert_before(new_article)
        first_existing.insert_before('\n')
    else:
        target.append(new_article)
    for duplicate_board in soup.select('section.article-index-board'):
        duplicate_board.decompose()
    final_html = doctype + '\n' + str(soup)
    output_html_path.write_text(final_html, encoding='utf-8')
    return output_html_path


def update_index_html(input_html_path, output_html_path, post_id, title, preview, update_floating=True, update_slider=True, replace_duplicate=True, max_cards=6):
    html_text = input_html_path.read_text(encoding='utf-8')
    doctype = detect_doctype(html_text)
    soup = BeautifulSoup(html_text, 'html.parser')
    if update_floating:
        notif_text = soup.select_one('div.floating-notif div.notif-text')
        if notif_text is not None:
            notif_text.clear()
            strong = soup.new_tag('strong')
            strong.string = 'Research Watch:'
            notif_text.append(strong)
            notif_text.append(NavigableString(f' {title}. '))
            link = soup.new_tag('a', href=f'ongoing-work.html#{post_id}')
            link.string = 'Read note ->'
            notif_text.append(link)
        else:
            research_alert = soup.select_one('div.research-alert')
            if research_alert is not None:
                badge = research_alert.select_one('.research-alert-badge')
                if badge is not None:
                    badge.string = 'Latest note'
                chips = research_alert.select_one('.hero-chips')
                if chips is None:
                    chips = soup.new_tag('div', attrs={'class': 'hero-chips'})
                    research_alert.append(chips)
                chips.clear()
                chip_span = soup.new_tag('span')
                chip_span.string = preview or title
                chips.append(chip_span)
                link_node = research_alert.select_one('a.research-alert-link')
                if link_node is None:
                    link_node = soup.new_tag('a', attrs={'class': 'research-alert-link'})
                    research_alert.append(link_node)
                link_node['href'] = f'ongoing-work.html#{post_id}'
                link_node.string = 'Read the research watch ->'
            else:
                raise ValueError('Could not find a supported floating/latest-note block in index.html (expected either .floating-notif .notif-text or .research-alert).')
    if update_slider:
        watch_track = soup.select_one('div#watchTrack')
        if watch_track is not None:
            if replace_duplicate:
                for a in watch_track.select('a.watch-card'):
                    href = a.get('href', '')
                    if href.endswith(f'#{post_id}'):
                        a.decompose()
            card_fragment = BeautifulSoup(build_home_watch_card(post_id, title, preview), 'html.parser')
            new_card = card_fragment.find('a')
            first_card = watch_track.find('a', class_='watch-card')
            if first_card:
                first_card.insert_before('\n')
                first_card.insert_before(new_card)
                first_card.insert_before('\n')
            else:
                watch_track.append(new_card)
            cards = watch_track.select('a.watch-card')
            if max_cards and len(cards) > max_cards:
                for card in cards[max_cards:]:
                    card.decompose()
    for duplicate_board in soup.select('section.article-index-board'):
        duplicate_board.decompose()
    final_html = doctype + '\n' + str(soup)
    output_html_path.write_text(final_html, encoding='utf-8')
    return output_html_path


def write_full_post_html(output_path: Path, full_post_html: str):
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(full_post_html, encoding='utf-8')
    return output_path


if BATCH_PROCESS_ALL_DOCX:
    DOCX_PATHS = list_processable_docx_files(WEEKLY_INPUTS_DIR)
else:
    DOCX_PATHS = [pick_latest_docx(WEEKLY_INPUTS_DIR)] if AUTO_PICK_LATEST_DOCX else [MANUAL_DOCX_PATH]
LATEST_DOCX_PATH = DOCX_PATHS[-1]
print('DOCX files selected:')
for p in DOCX_PATHS:
    latest_flag = '   <-- candidate for homepage latest note' if p == LATEST_DOCX_PATH else ''
    print(f'- {p}{latest_flag}')
path_warnings = validate_paths(LATEST_DOCX_PATH, INPUT_ONGOING_HTML, OUTPUT_ONGOING_HTML, INPUT_INDEX_HTML, OUTPUT_INDEX_HTML)
if path_warnings:
    print('CONFIG WARNINGS:')
    for w in path_warnings:
        print('-', w)
    print()
processed_items = []
for docx_path in DOCX_PATHS:
    docx_stem = docx_path.stem
    relative_full_post_link = f'posts/{docx_stem}.html'
    suggested_post_html_path = OUTPUT_POSTS_DIR / f'{docx_stem}.html'
    data = parse_weekly_docx(docx_path)
    if is_placeholder_docx(data):
        print(f'SKIPPING TEMPLATE-LIKE DOCX: {docx_path.name}')
        print('- This file still contains template placeholder text/links, so it will not be published or used for homepage updates.\n')
        continue
    normalize_warnings = normalize_docx_data(data, docx_stem, relative_full_post_link)
    md = data['metadata']
    if normalize_warnings:
        print(f'NORMALIZATION NOTES for {docx_path.name}:')
        for w in normalize_warnings:
            print('-', w)
        print()
    processed_items.append({'docx_path': docx_path, 'docx_stem': docx_stem, 'suggested_post_html_path': suggested_post_html_path, 'data': data, 'post_id': _clean(md.get('Post ID')), 'title': _clean(md.get('Title')), 'meta_line': _clean(md.get('Meta Line')), 'preview': _clean(md.get('Preview')), 'category': _clean(md.get('Category')), 'full_post_link': _clean(md.get('Full Post Link (optional)'))})

if not processed_items:
    raise RuntimeError('No publishable DOCX files were found after filtering out template-like placeholders.')
print('Selected DOCX summary:')
for item in processed_items:
    latest_tag = ' [homepage latest note]' if item['docx_path'] == LATEST_DOCX_PATH else ''
    print(f"- {item['docx_path'].name}{latest_tag}")
    print('  Post ID:', item['post_id'])
    print('  Category:', item['category'])
    print('  Title:', item['title'])
    print('  Meta Line:', item['meta_line'])
    print('  Full post link:', item['full_post_link'])
    print('  Generated full post page:', item['suggested_post_html_path'])
    print()
ongoing_output_path = None
index_output_path = None
latest_home_item = None
generated_post_paths = []
publishable_latest_docx = max((item['docx_path'] for item in processed_items), key=_docx_sort_key)
for item in processed_items:
    post_id, title, preview, category, article_html = build_watch_article(item['data'])
    ongoing_input_for_this_run = OUTPUT_ONGOING_HTML if ongoing_output_path is not None else INPUT_ONGOING_HTML
    ongoing_output_path = update_ongoing_work_html(ongoing_input_for_this_run, OUTPUT_ONGOING_HTML, post_id, article_html, category, REPLACE_DUPLICATE_POST_ID)
    index_input_for_this_run = OUTPUT_INDEX_HTML if index_output_path is not None else INPUT_INDEX_HTML
    index_output_path = update_index_html(index_input_for_this_run, OUTPUT_INDEX_HTML, post_id, title, preview, False, UPDATE_HOME_SLIDER, REPLACE_DUPLICATE_POST_ID, HOME_SLIDER_MAX_CARDS)
    if AUTO_GENERATE_FULL_POST:
        full_post_html = build_full_post_html(item['data'])
        generated_post_path = write_full_post_html(item['suggested_post_html_path'], full_post_html)
        generated_post_paths.append(generated_post_path)
    if item['docx_path'] == publishable_latest_docx:
        latest_home_item = {'post_id': post_id, 'title': title, 'preview': preview}
if latest_home_item is None:
    raise RuntimeError('Could not determine which processed DOCX should control the homepage latest note.')
index_input_for_floating = OUTPUT_INDEX_HTML if index_output_path is not None else INPUT_INDEX_HTML
index_output_path = update_index_html(index_input_for_floating, OUTPUT_INDEX_HTML, latest_home_item['post_id'], latest_home_item['title'], latest_home_item['preview'], UPDATE_HOME_FLOATING, False, REPLACE_DUPLICATE_POST_ID, HOME_SLIDER_MAX_CARDS)
print('Updated ongoing-work.html:', ongoing_output_path.resolve())
print('Updated index.html:', index_output_path.resolve())
if generated_post_paths:
    print('Generated/updated full posts:')
    for p in generated_post_paths:
        print('-', p.resolve())
print()
print('Homepage latest note now points to:', latest_home_item['post_id'])
print('Homepage latest note title:', latest_home_item['title'])
